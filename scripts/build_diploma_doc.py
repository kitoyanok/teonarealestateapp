from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "Пояснительная_записка_Тэона.md"
OUTPUT = ROOT / "Пояснительная_записка_Тэона.docx"


def apply_base_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)

    for style_name, size in [("Title", 16), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 14)]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True


def set_paragraph_format(paragraph, center: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = Pt(6)
    fmt.first_line_indent = Pt(35) if not center else Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph(style=f"Heading {min(level, 3)}")
    run = paragraph.add_run(text.strip())
    run.bold = True
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.style = document.styles["List Bullet"]
    paragraph.add_run(text.strip())
    set_paragraph_format(paragraph)


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.add_run(text.strip())
    set_paragraph_format(paragraph)


def add_blank(document: Document) -> None:
    paragraph = document.add_paragraph("")
    paragraph.paragraph_format.space_after = Pt(0)


def build() -> None:
    document = Document()
    for section in document.sections:
      section.top_margin = Pt(72)
      section.bottom_margin = Pt(72)
      section.left_margin = Pt(72)
      section.right_margin = Pt(72)

    apply_base_styles(document)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    for raw in lines:
        line = raw.rstrip()
        if not line:
            add_blank(document)
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph(style="Title")
            paragraph.add_run(line[2:].strip())
            set_paragraph_format(paragraph, center=True)
            continue
        if line.startswith("## "):
            add_heading(document, line[3:], 1)
            continue
        if line.startswith("### "):
            add_heading(document, line[4:], 2)
            continue
        if line.startswith("- "):
            add_bullet(document, line[2:])
            continue
        if line[0].isdigit() and ". " in line[:4]:
            paragraph = document.add_paragraph(style="List Number")
            paragraph.add_run(line.split(". ", 1)[1].strip())
            set_paragraph_format(paragraph)
            continue
        add_paragraph(document, line)

    document.save(OUTPUT)


if __name__ == "__main__":
    build()
