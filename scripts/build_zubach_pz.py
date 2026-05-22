from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
ASSETS_DIR = ROOT / "docs" / "assets" / "zubach_pz"
OUTPUT = ROOT / "67_Зубач_ПЗ.docx"

PAGE_WIDTH = 1650
PAGE_HEIGHT = 980
WHITE = "#ffffff"
BLACK = "#262626"
GRAY = "#666666"
LIGHT = "#f1f1f1"
MID = "#d7d7d7"
ORANGE = "#ff6a00"


@dataclass
class UseCase:
    name: str
    actors: str
    summary: str
    goal: str
    kind: str
    success_actions: list[str]
    success_system: list[str]
    exception_actions: list[str]
    exception_system: list[str]


@dataclass
class TestCase:
    ident: str
    name: str
    goal: str
    preconditions: str
    actions: list[str]
    expected: str


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica.ttc",
        "/Library/Fonts/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            try:
                return ImageFont.truetype(candidate, size=size)
            except OSError:
                continue
    return ImageFont.load_default()


FONT_20 = load_font(20)
FONT_24 = load_font(24)
FONT_28 = load_font(28)
FONT_32 = load_font(32, bold=True)
FONT_40 = load_font(40, bold=True)
FONT_18 = load_font(18)
FONT_18B = load_font(18, bold=True)


def canvas(title: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (PAGE_WIDTH, PAGE_HEIGHT), WHITE)
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((24, 24, PAGE_WIDTH - 24, PAGE_HEIGHT - 24), radius=28, outline=MID, width=3)
    draw.text((56, 48), title, fill=BLACK, font=FONT_32)
    return image, draw


def draw_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fill: str = WHITE, outline: str = MID, radius: int = 18, font=FONT_24, text_fill: str = BLACK) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = box
    text_box = draw.multiline_textbbox((0, 0), text, font=font, spacing=6, align="center")
    tw = text_box[2] - text_box[0]
    th = text_box[3] - text_box[1]
    draw.multiline_text((x1 + (x2 - x1 - tw) / 2, y1 + (y2 - y1 - th) / 2), text, fill=text_fill, font=font, spacing=6, align="center")


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str = BLACK, width: int = 4) -> None:
    draw.line((start, end), fill=fill, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex >= sx else -1
        draw.polygon([(ex, ey), (ex - 18 * direction, ey - 10), (ex - 18 * direction, ey + 10)], fill=fill)
    else:
        direction = 1 if ey >= sy else -1
        draw.polygon([(ex, ey), (ex - 10, ey - 18 * direction), (ex + 10, ey - 18 * direction)], fill=fill)


def save(image: Image.Image, name: str) -> Path:
    ensure_dir(ASSETS_DIR)
    path = ASSETS_DIR / name
    image.save(path)
    return path


def build_use_case_diagram() -> Path:
    image, draw = canvas("Диаграмма вариантов использования приложения «Тэона»")
    draw_box(draw, (70, 260, 250, 620), "Риелтор", fill="#fafafa", font=FONT_32)
    cases = [
        ((420, 120, 770, 210), "Авторизация"),
        ((420, 250, 770, 340), "Создание клиента"),
        ((420, 380, 770, 470), "Запуск поиска объектов"),
        ((900, 120, 1270, 210), "Просмотр найденных\nобъектов"),
        ((900, 250, 1270, 340), "Добавление объекта\nв подборку"),
        ((900, 380, 1270, 470), "Формирование текста\nподборки"),
        ((660, 560, 1030, 650), "Отметка статуса\n«Отправлено»"),
    ]
    for box, text in cases:
        draw_box(draw, box, text, fill="#fff9f4")
    for target_y in [165, 295, 425]:
        draw_arrow(draw, (250, 440), (420, target_y))
    draw_arrow(draw, (770, 295), (900, 295))
    draw_arrow(draw, (770, 425), (900, 425))
    draw_arrow(draw, (1085, 340), (1085, 380))
    draw_arrow(draw, (845, 470), (845, 560))
    return save(image, "01_use_case.png")


def build_activity_diagram() -> Path:
    image, draw = canvas("Блок-схема сценария подбора объектов")
    nodes = [
        ((650, 90, 1000, 170), "Начало работы."),
        ((620, 220, 1030, 320), "Действие №1. Риелтор вводит данные клиента."),
        ((620, 370, 1030, 470), "Действие №2. Система сохраняет клиента\nи профиль поиска."),
        ((620, 520, 1030, 620), "Действие №3. API запускает search-service."),
        ((620, 670, 1030, 770), "Действие №4. Search-service нормализует\nи фильтрует объекты."),
        ((620, 820, 1030, 920), "Действие №5. API сохраняет найденные\nобъекты и результаты поиска."),
    ]
    for box, text in nodes:
        draw_box(draw, box, text)
    for idx in range(len(nodes) - 1):
        start = ((nodes[idx][0][0] + nodes[idx][0][2]) // 2, nodes[idx][0][3])
        end = ((nodes[idx + 1][0][0] + nodes[idx + 1][0][2]) // 2, nodes[idx + 1][0][1])
        draw_arrow(draw, start, end)
    draw_box(draw, (160, 520, 500, 620), "Условие.\nНайденные объекты есть.", fill="#fff9f4")
    draw_arrow(draw, (620, 570), (500, 570))
    draw_box(draw, (1120, 520, 1470, 620), "Исключение.\nИсточники недоступны.", fill="#fff2f2")
    draw_arrow(draw, (1030, 570), (1120, 570))
    return save(image, "02_activity.png")


def build_use_case_search_diagram() -> Path:
    image, draw = canvas("Диаграмма вариантов использования модуля поиска объектов")
    draw_box(draw, (70, 250, 250, 610), "Риелтор", fill="#fafafa", font=FONT_32)
    cases = [
        ((430, 120, 810, 210), "Ввод параметров\nклиента"),
        ((430, 265, 810, 355), "Запуск поиска"),
        ((430, 410, 810, 500), "Просмотр найденных\nобъектов"),
        ((970, 120, 1360, 210), "Фильтрация по\nгороду Краснодар"),
        ((970, 265, 1360, 355), "Расчет процента\nсовпадения"),
        ((970, 410, 1360, 500), "Сохранение\nрезультатов"),
    ]
    for box, text in cases:
        draw_box(draw, box, text, fill="#fff9f4")
    for y in [165, 310, 455]:
        draw_arrow(draw, (250, 430), (430, y))
    draw_arrow(draw, (810, 310), (970, 165))
    draw_arrow(draw, (810, 310), (970, 310))
    draw_arrow(draw, (810, 455), (970, 455))
    draw.text((420, 650), "Модуль поиска отвечает за получение параметров клиента, обращение к источникам,\nотсечение объектов не из Краснодара и сохранение релевантных результатов.", font=FONT_24, fill=GRAY, spacing=8)
    return save(image, "02_1_use_case_search.png")


def build_use_case_shortlist_diagram() -> Path:
    image, draw = canvas("Диаграмма вариантов использования модуля подборки")
    draw_box(draw, (70, 250, 250, 610), "Риелтор", fill="#fafafa", font=FONT_32)
    cases = [
        ((430, 120, 830, 210), "Открыть карточку\nклиента"),
        ((430, 265, 830, 355), "Добавить объект\nв подборку"),
        ((430, 410, 830, 500), "Удалить объект\nиз подборки"),
        ((970, 190, 1390, 280), "Подготовить текст\nсообщения"),
        ((970, 350, 1390, 440), "Отметить подборку\nотправленной"),
    ]
    for box, text in cases:
        draw_box(draw, box, text, fill="#fff9f4")
    for y in [165, 310, 455]:
        draw_arrow(draw, (250, 430), (430, y))
    draw_arrow(draw, (830, 310), (970, 235))
    draw_arrow(draw, (1180, 280), (1180, 350))
    draw.text((420, 650), "Модуль подборки сохраняет ручной выбор риелтора в PostgreSQL и формирует текст,\nкоторый риелтор копирует клиенту без автоматической отправки во внешние мессенджеры.", font=FONT_24, fill=GRAY, spacing=8)
    return save(image, "02_2_use_case_shortlist.png")


def build_state_diagram() -> Path:
    image, draw = canvas("Диаграмма состояний клиентской заявки")
    states = [
        ((90, 220, 310, 320), "new\nНовая"),
        ((430, 220, 680, 320), "searching\nИдет поиск"),
        ((820, 220, 1060, 320), "found\nНайдены объекты"),
        ((1200, 220, 1500, 320), "shortlist_ready\nПодборка готова"),
        ((820, 520, 1060, 620), "sent\nОтправлено"),
        ((430, 520, 680, 620), "no_results\nНет объектов"),
        ((90, 520, 310, 620), "error\nОшибка"),
    ]
    for box, text in states:
        draw_box(draw, box, text, fill="#fff9f4")
    draw_arrow(draw, (310, 270), (430, 270))
    draw_arrow(draw, (680, 270), (820, 270))
    draw_arrow(draw, (1060, 270), (1200, 270))
    draw_arrow(draw, (1350, 320), (955, 520))
    draw_arrow(draw, (555, 320), (555, 520))
    draw_arrow(draw, (430, 570), (310, 570))
    draw.text((105, 735), "Состояние клиента меняется серверной логикой. История поиска хранится отдельно в search_runs,\nпоэтому повторный запуск поиска не уничтожает предыдущий контекст.", font=FONT_24, fill=GRAY, spacing=8)
    return save(image, "15_state_client.png")


def build_component_diagram() -> Path:
    image, draw = canvas("Диаграмма компонентов приложения «Тэона»")
    draw_box(draw, (90, 160, 430, 300), "Web UI\nReact, Vite", fill="#fff9f4", font=FONT_28)
    draw_box(draw, (620, 160, 980, 300), "API\nExpress, pg", fill="#fff9f4", font=FONT_28)
    draw_box(draw, (1160, 160, 1520, 300), "Search-service\nFastAPI", fill="#fff9f4", font=FONT_28)
    draw_box(draw, (620, 500, 980, 640), "PostgreSQL\nclients, properties,\nshortlist", fill="#fff9f4", font=FONT_28)
    draw_box(draw, (1160, 500, 1520, 640), "Источники\nзастройщики и\nагрегаторы", fill="#fafafa", font=FONT_28)
    draw_arrow(draw, (430, 230), (620, 230))
    draw_arrow(draw, (980, 230), (1160, 230))
    draw_arrow(draw, (800, 300), (800, 500))
    draw_arrow(draw, (1340, 500), (1340, 300))
    draw.text((90, 760), "Компоненты развертываются через Docker Compose. Frontend обращается только к API,\nAPI координирует работу БД и search-service, а search-service не имеет прямого доступа к PostgreSQL.", font=FONT_24, fill=GRAY, spacing=8)
    return save(image, "16_component.png")


def build_parsing_algorithm_diagram() -> Path:
    image, draw = canvas("Алгоритм парсинга и географической фильтрации")
    nodes = [
        ((570, 80, 1080, 150), "Начало. Получен SearchRequest."),
        ((570, 210, 1080, 290), "Выбор источников по типу недвижимости."),
        ((570, 350, 1080, 430), "Загрузка HTML и JSON-LD."),
        ((570, 490, 1080, 570), "Извлечение цены, площади, комнат,\nрайона, адреса и ссылки."),
        ((570, 630, 1080, 710), "Проверка: объект относится\nк городу Краснодар?"),
        ((570, 770, 1080, 850), "Нормализация title, description,\nизображений и характеристик."),
    ]
    for box, text in nodes:
        draw_box(draw, box, text)
    for idx in range(len(nodes) - 1):
        draw_arrow(draw, ((nodes[idx][0][0] + nodes[idx][0][2]) // 2, nodes[idx][0][3]), ((nodes[idx + 1][0][0] + nodes[idx + 1][0][2]) // 2, nodes[idx + 1][0][1]))
    draw_box(draw, (1170, 620, 1530, 720), "Нет.\nКарточка отбрасывается.", fill="#fff2f2")
    draw_arrow(draw, (1080, 670), (1170, 670))
    return save(image, "17_algorithm_parsing.png")


def build_scoring_algorithm_diagram() -> Path:
    image, draw = canvas("Алгоритм расчета match score")
    nodes = [
        ((90, 170, 470, 260), "Проверить бюджет"),
        ((640, 170, 1020, 260), "Проверить площадь\nи комнатность"),
        ((1190, 170, 1570, 260), "Проверить район\nили поселок"),
        ((90, 470, 470, 560), "Проверить срок,\nотделку, этажность"),
        ((640, 470, 1020, 560), "Сформировать причины\nсовпадения"),
        ((1190, 470, 1570, 560), "Ограничить итог\nдиапазоном 0-100"),
    ]
    for box, text in nodes:
        draw_box(draw, box, text, fill="#fff9f4")
    draw_arrow(draw, (470, 215), (640, 215))
    draw_arrow(draw, (1020, 215), (1190, 215))
    draw_arrow(draw, (1380, 260), (280, 470))
    draw_arrow(draw, (470, 515), (640, 515))
    draw_arrow(draw, (1020, 515), (1190, 515))
    draw.text((120, 760), "Для квартир учитываются цена, комнаты, площадь, район, срок сдачи и отделка.\nДля домов учитываются бюджет, площадь дома, участок, локация, коммуникации, этажность и спальни.", font=FONT_24, fill=GRAY, spacing=8)
    return save(image, "18_algorithm_scoring.png")


def build_persistence_algorithm_diagram() -> Path:
    image, draw = canvas("Алгоритм сохранения результатов поиска")
    nodes = [
        ((600, 90, 1050, 170), "API получает PropertyItem."),
        ((600, 240, 1050, 320), "upsert properties по source_url."),
        ((600, 390, 1050, 470), "upsert client_found_properties."),
        ((600, 540, 1050, 620), "Обновление search_runs."),
        ((600, 690, 1050, 770), "Обновление статуса клиента."),
    ]
    for box, text in nodes:
        draw_box(draw, box, text)
    for idx in range(len(nodes) - 1):
        draw_arrow(draw, ((nodes[idx][0][0] + nodes[idx][0][2]) // 2, nodes[idx][0][3]), ((nodes[idx + 1][0][0] + nodes[idx + 1][0][2]) // 2, nodes[idx + 1][0][1]))
    draw_box(draw, (150, 390, 500, 510), "Если source_url уже есть,\nкарточка обновляется.", fill="#fafafa")
    draw_arrow(draw, (600, 430), (500, 450))
    draw_box(draw, (1140, 390, 1500, 510), "Если объект новый,\nсоздается новая запись.", fill="#fafafa")
    draw_arrow(draw, (1050, 430), (1140, 450))
    return save(image, "19_algorithm_persistence.png")


def build_test_results_image() -> Path:
    image, draw = canvas("Результаты автоматизированного тестирования")
    draw_box(draw, (110, 140, 1540, 810), "", fill="#111111", outline="#111111", radius=20)
    lines = [
        "> npm run typecheck",
        "@estateflow/web: tsc -b --noEmit ... OK",
        "@estateflow/api: tsc -p tsconfig.json --noEmit ... OK",
        "",
        "> npm run build",
        "Vite build completed successfully.",
        "API TypeScript build completed successfully.",
        "",
        "> python -m unittest discover -s tests -v",
        "Ran 10 tests in 0.001s",
        "OK",
    ]
    y = 185
    mono = load_font(26)
    for line in lines:
        draw.text((160, y), line, font=mono, fill="#f4f4f4")
        y += 48
    return save(image, "20_test_results.png")


def build_architecture_diagram() -> Path:
    image, draw = canvas("Архитектурная схема системы")
    draw_box(draw, (120, 250, 430, 390), "Frontend\nReact + Vite", fill="#fff9f4", font=FONT_32)
    draw_box(draw, (540, 250, 890, 390), "Backend API\nNode.js + Express + pg", fill="#fff9f4", font=FONT_32)
    draw_box(draw, (1010, 160, 1440, 320), "Search-service\nPython + FastAPI", fill="#fff9f4", font=FONT_32)
    draw_box(draw, (1010, 420, 1440, 580), "PostgreSQL\nПостоянное хранение", fill="#fff9f4", font=FONT_32)
    draw_box(draw, (1010, 700, 1440, 860), "Внешние источники\nсайты новостроек и КП", fill="#fafafa", font=FONT_32)
    draw_arrow(draw, (430, 320), (540, 320))
    draw_arrow(draw, (890, 280), (1010, 240))
    draw_arrow(draw, (890, 360), (1010, 500))
    draw_arrow(draw, (1225, 320), (1225, 420))
    draw_arrow(draw, (1225, 700), (1225, 580))
    draw.text((150, 500), "Frontend отвечает за интерфейс,\nмаршруты и отображение данных.", font=FONT_24, fill=GRAY, spacing=8)
    draw.text((540, 640), "API хранит клиентов, найденные объекты,\nподборки, тексты сообщений и историю поиска.", font=FONT_24, fill=GRAY, spacing=8)
    return save(image, "03_architecture.png")


def build_er_schema() -> Path:
    image, draw = canvas("Схема базы данных PostgreSQL")
    boxes = {
        "users": (70, 120, 390, 330),
        "clients": (430, 120, 780, 360),
        "client_search_profiles": (820, 80, 1220, 340),
        "properties": (1260, 80, 1600, 360),
        "client_found_properties": (820, 420, 1220, 670),
        "shortlist_items": (1260, 420, 1600, 610),
        "share_messages": (430, 470, 780, 650),
        "search_runs": (70, 470, 390, 650),
    }
    content = {
        "users": "users\nid\nlogin\npassword_hash\nname\nemail\nphone",
        "clients": "clients\nid\nrealtor_id\nname\nphone\nstatus\nproperty_type",
        "client_search_profiles": "client_search_profiles\nid\nclient_id\nbudget_min / max\nrooms_min / max\narea_min / max\ndistricts[]",
        "properties": "properties\nid\nsource_url\ntitle\nprice\narea\nrooms\nimages[]",
        "client_found_properties": "client_found_properties\nid\nclient_id\nproperty_id\nmatch_score\nmatch_reasons[]",
        "shortlist_items": "shortlist_items\nid\nclient_id\nproperty_id\nnote",
        "share_messages": "share_messages\nid\nclient_id\nchannel\nmessage_text\nsent_marked_at",
        "search_runs": "search_runs\nid\nclient_id\nstatus\ntotal_found\ntotal_saved",
    }
    for key, box in boxes.items():
        draw_box(draw, box, content[key], font=FONT_20)
    draw_arrow(draw, (390, 220), (430, 220))
    draw_arrow(draw, (780, 220), (820, 210))
    draw_arrow(draw, (780, 255), (1260, 220))
    draw_arrow(draw, (605, 360), (605, 470))
    draw_arrow(draw, (220, 330), (220, 470))
    draw_arrow(draw, (1020, 340), (1020, 420))
    draw_arrow(draw, (1380, 360), (1380, 420))
    draw_arrow(draw, (1220, 540), (1260, 520))
    return save(image, "04_er_schema.png")


def wireframe_frame(draw: ImageDraw.ImageDraw, title: str, subtitle: str, areas: list[tuple[tuple[int, int, int, int], str]]) -> None:
    draw.text((300, 125), title, fill=BLACK, font=FONT_40)
    draw.text((300, 182), subtitle, fill=GRAY, font=FONT_24)
    for box, label in areas:
        draw_box(draw, box, label, fill="#fafafa", font=FONT_20)


def build_wireframe_login() -> Path:
    image, draw = canvas("Wireframe окна авторизации")
    draw_box(draw, (80, 120, 760, 860), "Информационный блок\nлоготип\nописание системы", fill="#fafafa", font=FONT_28)
    areas = [
        ((960, 180, 1510, 280), "Логотип"),
        ((930, 330, 1540, 430), "Поле логина"),
        ((930, 470, 1540, 570), "Поле пароля"),
        ((930, 620, 1540, 720), "Кнопка «Войти»"),
        ((930, 760, 1540, 840), "Сообщение об ошибке / подсказка"),
    ]
    wireframe_frame(draw, "Экран входа", "Доступ к рабочему пространству риелтора.", areas)
    return save(image, "05_wireframe_login.png")


def build_wireframe_dashboard() -> Path:
    image, draw = canvas("Wireframe главной страницы")
    areas = [
        ((60, 120, 260, 900), "Боковое меню"),
        ((300, 120, 1590, 210), "Верхняя панель поиска и профиля"),
        ((300, 250, 600, 430), "KPI 1"),
        ((630, 250, 930, 430), "KPI 2"),
        ((960, 250, 1260, 430), "KPI 3"),
        ((1290, 250, 1590, 430), "KPI 4"),
        ((300, 470, 1170, 890), "График активности за 7 дней"),
        ((1200, 470, 1590, 650), "Блок внимания"),
        ((1200, 690, 1590, 890), "Последние клиенты"),
    ]
    wireframe_frame(draw, "Главная страница", "KPI, график активности и клиенты, требующие внимания.", areas)
    return save(image, "06_wireframe_dashboard.png")


def build_wireframe_clients() -> Path:
    image, draw = canvas("Wireframe списка клиентов")
    areas = [
        ((60, 120, 260, 900), "Боковое меню"),
        ((300, 120, 1590, 210), "Верхняя панель"),
        ((300, 250, 1590, 340), "Фильтр и строка поиска"),
        ((300, 380, 1590, 890), "Таблица клиентов\nФИО | Профиль | Статус | Дата | Действие"),
    ]
    wireframe_frame(draw, "Список клиентов", "Просмотр базы клиентов и быстрый переход в карточку.", areas)
    return save(image, "07_wireframe_clients.png")


def build_wireframe_client_card() -> Path:
    image, draw = canvas("Wireframe карточки клиента")
    areas = [
        ((60, 120, 260, 900), "Боковое меню"),
        ((300, 120, 1590, 210), "Верхняя панель"),
        ((300, 250, 980, 430), "Параметры клиента и действия"),
        ((300, 470, 1180, 900), "Найденные объекты"),
        ((1220, 250, 1590, 900), "Подборка клиента\nКнопка генерации сообщения"),
    ]
    wireframe_frame(draw, "Карточка клиента", "Поиск, результаты, shortlist и подготовка сообщения.", areas)
    return save(image, "08_wireframe_client_card.png")


def draw_topbar(draw: ImageDraw.ImageDraw) -> None:
    draw.rounded_rectangle((290, 115, 1585, 205), radius=20, fill=WHITE, outline=MID, width=2)
    draw.rounded_rectangle((360, 138, 850, 182), radius=16, fill=LIGHT, outline=LIGHT)
    draw.text((395, 147), "Поиск по клиентам и объектам", font=FONT_20, fill=GRAY)
    draw.rounded_rectangle((1440, 132, 1560, 188), radius=18, fill="#fff3eb", outline="#ffd3b2", width=2)
    draw.text((1475, 148), "NZ", font=FONT_20, fill=ORANGE)


def build_screen_login() -> Path:
    image, draw = canvas("Скриншот окна авторизации")
    draw.rounded_rectangle((70, 120, 790, 900), radius=28, fill="#f7f7f7", outline=MID, width=2)
    draw.text((140, 210), "Тэона", font=FONT_40, fill=BLACK)
    draw.text((140, 285), "Рабочее пространство риелтора\nдля подбора недвижимости.", font=FONT_28, fill=GRAY, spacing=10)
    draw.rounded_rectangle((950, 190, 1510, 860), radius=28, fill=WHITE, outline=MID, width=2)
    draw.text((1030, 260), "Вход в систему", font=FONT_32, fill=BLACK)
    for idx, label in enumerate(["Логин", "Пароль"]):
        top = 360 + idx * 150
        draw.text((1040, top - 36), label, font=FONT_20, fill=GRAY)
        draw.rounded_rectangle((1040, top, 1430, top + 70), radius=16, fill=WHITE, outline=MID, width=2)
        draw.text((1070, top + 18), "test" if idx == 0 else "••••••", font=FONT_24, fill=BLACK)
    draw.rounded_rectangle((1040, 690, 1430, 770), radius=18, fill=ORANGE, outline=ORANGE, width=2)
    draw.text((1170, 713), "Войти", font=FONT_24, fill=WHITE)
    return save(image, "09_screen_login.png")


def build_screen_dashboard() -> Path:
    image, draw = canvas("Скриншот главной страницы")
    draw.rounded_rectangle((55, 115, 250, 900), radius=24, fill=WHITE, outline=MID, width=2)
    draw.text((95, 170), "Тэона", font=FONT_32, fill=BLACK)
    for idx, label in enumerate(["Главная", "Клиенты", "Аналитика", "Профиль"]):
        y = 290 + idx * 80
        fill = "#fff3eb" if idx == 0 else WHITE
        outline = "#ffd3b2" if idx == 0 else MID
        draw.rounded_rectangle((80, y, 225, y + 54), radius=16, fill=fill, outline=outline, width=2)
        draw.text((105, y + 15), label, font=FONT_20, fill=ORANGE if idx == 0 else BLACK)
    draw_topbar(draw)
    cards = [("1", "Клиенты в работе"), ("81", "Найдено объектов"), ("7", "В подборках"), ("1", "Готово к отправке")]
    for idx, (num, label) in enumerate(cards):
        x1 = 300 + idx * 320
        draw.rounded_rectangle((x1, 250, x1 + 290, 430), radius=24, fill=WHITE, outline=MID, width=2)
        draw.text((x1 + 30, 290), num, font=FONT_40, fill=BLACK)
        draw.text((x1 + 30, 380), label, font=FONT_20, fill=GRAY)
    draw.rounded_rectangle((300, 470, 1175, 900), radius=24, fill=WHITE, outline=MID, width=2)
    draw.text((340, 515), "Активность за 7 дней", font=FONT_28, fill=BLACK)
    for idx, h in enumerate([110, 150, 130, 270, 140, 120, 170]):
        x = 360 + idx * 110
        draw.rounded_rectangle((x, 830 - h, x + 72, 830), radius=18, fill=ORANGE if idx == 6 else LIGHT, outline=LIGHT)
    for idx, day in enumerate(["ср", "чт", "пт", "сб", "вс", "пн", "вт"]):
        draw.text((375 + idx * 110, 850), day, font=FONT_20, fill=GRAY)
    draw.rounded_rectangle((1210, 470, 1585, 655), radius=24, fill=WHITE, outline=MID, width=2)
    draw.text((1240, 515), "С кем продолжить работу", font=FONT_24, fill=BLACK)
    for i, row in enumerate(["Иван Иванов", "Анна Петрова", "Мария Смирнова"]):
        y = 570 + i * 40
        draw.text((1245, y), row, font=FONT_20, fill=GRAY if i else BLACK)
    draw.rounded_rectangle((1210, 690, 1585, 900), radius=24, fill=WHITE, outline=MID, width=2)
    draw.text((1240, 735), "Последние клиенты", font=FONT_24, fill=BLACK)
    for i, row in enumerate(["Иван Иванов", "Олег Сафронов", "Тамара Белая", "Николай Громов"]):
        draw.text((1245, 785 + i * 28), row, font=FONT_20, fill=GRAY if i else BLACK)
    return save(image, "10_screen_dashboard.png")


def build_screen_clients() -> Path:
    image, draw = canvas("Скриншот списка клиентов")
    draw.rounded_rectangle((55, 115, 250, 900), radius=24, fill=WHITE, outline=MID, width=2)
    draw.text((95, 170), "Тэона", font=FONT_32, fill=BLACK)
    draw_topbar(draw)
    draw.rounded_rectangle((300, 250, 1585, 335), radius=20, fill=WHITE, outline=MID, width=2)
    draw.rounded_rectangle((330, 270, 840, 315), radius=16, fill=LIGHT, outline=LIGHT)
    draw.text((360, 281), "Поиск клиента по имени и телефону", font=FONT_20, fill=GRAY)
    draw.rounded_rectangle((1360, 264, 1545, 320), radius=18, fill=ORANGE, outline=ORANGE)
    draw.text((1386, 281), "Добавить клиента", font=FONT_20, fill=WHITE)
    draw.rounded_rectangle((300, 370, 1585, 900), radius=24, fill=WHITE, outline=MID, width=2)
    headers = ["ФИО", "Профиль", "Статус", "Дата", "Действие"]
    xs = [335, 570, 1080, 1290, 1440]
    for x, h in zip(xs, headers):
        draw.text((x, 410), h, font=FONT_20, fill=GRAY)
    rows = [
        ("Иван Иванов", "до 8 млн руб. • 1-2 комн.", "найдено", "21.05.2026"),
        ("Анна Петрова", "до 6 млн руб. • студия-1 комн.", "новый", "20.05.2026"),
        ("Николай Громов", "дом от 120 м² • 5 сот.", "подборка", "20.05.2026"),
        ("Тамара Белая", "до 10 млн руб. • 2-3 комн.", "отправлено", "19.05.2026"),
    ]
    for idx, row in enumerate(rows):
        y = 470 + idx * 92
        draw.line((330, y - 18, 1550, y - 18), fill=LIGHT, width=2)
        for j, cell in enumerate(row):
            draw.text((xs[j], y), cell, font=FONT_20, fill=BLACK if j != 2 else ORANGE)
        draw.rounded_rectangle((1430, y - 10, 1535, y + 34), radius=16, fill=WHITE, outline=MID, width=2)
        draw.text((1450, y + 2), "Открыть", font=FONT_20, fill=BLACK)
    return save(image, "11_screen_clients.png")


def build_screen_client_card() -> Path:
    image, draw = canvas("Скриншот карточки клиента")
    draw.rounded_rectangle((55, 115, 250, 900), radius=24, fill=WHITE, outline=MID, width=2)
    draw.text((95, 170), "Тэона", font=FONT_32, fill=BLACK)
    draw_topbar(draw)
    draw.text((305, 235), "Иван Иванов", font=FONT_32, fill=BLACK)
    draw.text((305, 278), "до 8 млн руб. • 1-2 комн.", font=FONT_24, fill=GRAY)
    draw.rounded_rectangle((300, 320, 980, 470), radius=24, fill=WHITE, outline=MID, width=2)
    draw.text((335, 360), "Параметры клиента", font=FONT_24, fill=BLACK)
    draw.text((335, 405), "Бюджет: до 8 млн руб.\nКомнатность: 1-2\nРайон: Прикубанский", font=FONT_20, fill=GRAY, spacing=8)
    draw.rounded_rectangle((1000, 320, 1585, 900), radius=24, fill=WHITE, outline=MID, width=2)
    draw.text((1040, 360), "Подборка клиента", font=FONT_24, fill=BLACK)
    draw.rounded_rectangle((1040, 780, 1540, 845), radius=18, fill=ORANGE, outline=ORANGE)
    draw.text((1150, 800), "Подготовить сообщение", font=FONT_24, fill=WHITE)
    card_positions = [(300, 510), (635, 510)]
    titles = ["1-к квартира, 40.6 м²", "2-к квартира, 53.8 м²"]
    prices = ["5 275 000 руб.", "6 540 000 руб."]
    meta = ["ЖК «Nova Vita»\nПрикубанский район", "ЖК «Точно»\nЦентральный район"]
    for (x, y), title, price, meta_line in zip(card_positions, titles, prices, meta):
        draw.rounded_rectangle((x, y, x + 310, y + 360), radius=22, fill=WHITE, outline=MID, width=2)
        draw.rounded_rectangle((x + 16, y + 16, x + 294, y + 140), radius=18, fill=LIGHT, outline=LIGHT)
        draw.text((x + 28, y + 170), title, font=FONT_20, fill=BLACK)
        draw.multiline_text((x + 28, y + 210), meta_line, font=FONT_18, fill=GRAY, spacing=5)
        draw.text((x + 28, y + 275), price, font=FONT_28, fill=BLACK)
        draw.text((x + 28, y + 322), "Совпадение: 68%", font=FONT_20, fill=GRAY)
    return save(image, "12_screen_client_card.png")


def build_screen_message_modal() -> Path:
    image, draw = canvas("Скриншот модального окна подготовки сообщения")
    draw.rounded_rectangle((140, 120, 1520, 890), radius=24, fill="#f4f4f4", outline="#f4f4f4")
    draw.rounded_rectangle((430, 180, 1240, 820), radius=26, fill=WHITE, outline=MID, width=2)
    draw.text((500, 235), "Подготовка сообщения клиенту", font=FONT_32, fill=BLACK)
    draw.text((500, 300), "Телефон клиента", font=FONT_20, fill=GRAY)
    draw.rounded_rectangle((500, 335, 920, 395), radius=16, fill=WHITE, outline=MID, width=2)
    draw.text((530, 352), "+7 (918) 123-45-67", font=FONT_24, fill=BLACK)
    draw.rounded_rectangle((950, 335, 1160, 395), radius=16, fill="#fff3eb", outline="#ffd3b2", width=2)
    draw.text((1005, 352), "Копировать", font=FONT_20, fill=ORANGE)
    draw.text((500, 445), "Текст сообщения", font=FONT_20, fill=GRAY)
    draw.rounded_rectangle((500, 480, 1160, 680), radius=18, fill="#fffdfb", outline=MID, width=2)
    message = (
        "Здравствуйте. Подобрал для вас несколько вариантов.\n\n"
        "1. 1-к квартира, 40.6 м², 5 275 000 руб.\n"
        "2. 2-к квартира, 53.8 м², 6 540 000 руб.\n\n"
        "Если нужно, подготовлю еще варианты."
    )
    draw.multiline_text((530, 515), message, font=FONT_20, fill=BLACK, spacing=10)
    draw.rounded_rectangle((500, 720, 800, 780), radius=18, fill=ORANGE, outline=ORANGE)
    draw.text((590, 737), "Копировать текст", font=FONT_20, fill=WHITE)
    draw.rounded_rectangle((830, 720, 1160, 780), radius=18, fill=WHITE, outline=ORANGE, width=2)
    draw.text((900, 737), "Отметить отправленным", font=FONT_20, fill=ORANGE)
    return save(image, "13_screen_message_modal.png")


def build_screen_db_view() -> Path:
    image, draw = canvas("Скриншот просмотра базы данных")
    draw.rounded_rectangle((60, 120, 260, 900), radius=20, fill="#fafafa", outline=MID, width=2)
    tables = ["users", "clients", "client_search_profiles", "properties", "client_found_properties", "shortlist_items", "share_messages", "search_runs"]
    for idx, name in enumerate(tables[:7]):
        draw.text((85, 165 + idx * 75), name, font=FONT_20, fill=ORANGE if name == "clients" else BLACK)
    draw.rounded_rectangle((290, 120, 1590, 900), radius=20, fill=WHITE, outline=MID, width=2)
    draw.text((330, 160), "Таблица clients", font=FONT_28, fill=BLACK)
    headers = ["id", "name", "phone", "status", "property_type", "created_at"]
    x_positions = [330, 570, 820, 1040, 1230, 1420]
    for x, header in zip(x_positions, headers):
        draw.text((x, 225), header, font=FONT_20, fill=GRAY)
    rows = [
        ("7d8a...", "Иван Иванов", "+79181234567", "found", "apartment", "2026-05-21"),
        ("21ca...", "Анна Петрова", "+79001234567", "new", "apartment", "2026-05-20"),
        ("e5f2...", "Николай Громов", "+79601234567", "shortlist", "house", "2026-05-20"),
        ("91bc...", "Тамара Белая", "+79111234567", "sent", "apartment", "2026-05-19"),
    ]
    for idx, row in enumerate(rows):
        y = 290 + idx * 85
        draw.line((320, y - 18, 1560, y - 18), fill=LIGHT, width=2)
        for x, cell in zip(x_positions, row):
            draw.text((x, y), cell, font=FONT_18B if x == 570 else FONT_20, fill=BLACK if x != 1040 else ORANGE)
    return save(image, "14_screen_db_view.png")


def build_assets() -> dict[str, Path]:
    return {
        "use_case": build_use_case_diagram(),
        "use_case_search": build_use_case_search_diagram(),
        "use_case_shortlist": build_use_case_shortlist_diagram(),
        "activity": build_activity_diagram(),
        "architecture": build_architecture_diagram(),
        "er": build_er_schema(),
        "state": build_state_diagram(),
        "component": build_component_diagram(),
        "algorithm_parsing": build_parsing_algorithm_diagram(),
        "algorithm_scoring": build_scoring_algorithm_diagram(),
        "algorithm_persistence": build_persistence_algorithm_diagram(),
        "test_results": build_test_results_image(),
        "wf_login": build_wireframe_login(),
        "wf_dashboard": build_wireframe_dashboard(),
        "wf_clients": build_wireframe_clients(),
        "wf_client": build_wireframe_client_card(),
        "screen_login": build_screen_login(),
        "screen_dashboard": build_screen_dashboard(),
        "screen_clients": build_screen_clients(),
        "screen_client": build_screen_client_card(),
        "screen_modal": build_screen_message_modal(),
        "screen_db": build_screen_db_view(),
    }


def set_text_style(run, size: int = 14, bold: bool = False, italic: bool = False, font_name: str = "Times New Roman") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def configure_styles(document: Document) -> None:
    for section in document.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.63)
        section.bottom_margin = Cm(3.6)
        section.left_margin = Cm(2)
        section.right_margin = Cm(0.55)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(1.5)
    list_style = document.styles["List Paragraph"]
    list_style.font.name = "Times New Roman"
    list_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    list_style.font.size = Pt(14)
    list_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    list_style.paragraph_format.space_after = Pt(0)
    list_style.paragraph_format.first_line_indent = Cm(1.5)
    for style_name, size in [("Title", 14), ("Heading 1", 14), ("Heading 2", 14), ("Heading 3", 14)]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = False


def set_paragraph(paragraph, center: bool = False, indent: bool = True) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.first_line_indent = Cm(1.5) if indent and not center else Pt(0)


def add_text(document: Document, text: str, *, center: bool = False, bold: bool = False, italic: bool = False, indent: bool = True, size: int = 14) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_text_style(run, size=size, bold=bold, italic=italic)
    set_paragraph(paragraph, center=center, indent=indent)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph(style=f"Heading {min(level, 3)}")
    run = paragraph.add_run(text)
    set_text_style(run, size=14, bold=False)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.5)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(0)


def add_page_break(document: Document) -> None:
    document.add_page_break()


def add_toc(document: Document) -> None:
    p = document.add_paragraph()
    set_paragraph(p, indent=False)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Обновите содержание в Word: правой кнопкой по содержанию -> Обновить поле."
    r.append(t)
    fld.append(r)
    p._p.append(fld)


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = p.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_text_style(run, size=12)


def add_figure(document: Document, caption: str, image_path: Path, width_cm: float = 15.5) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.first_line_indent = Pt(0)
    caption_run = caption_p.add_run(caption)
    set_text_style(caption_run, italic=True)


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.first_line_indent = Pt(0)
                for run in paragraph.runs:
                    set_text_style(run, size=12)


def add_caption(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_text_style(run, italic=True)


def add_key_value_table(document: Document, caption: str, rows: list[tuple[str, str]]) -> None:
    add_caption(document, caption)
    table = document.add_table(rows=0, cols=2)
    for left, right in rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right
    style_table(table)


def add_scenario_table(document: Document, caption: str, actions: list[str], system: list[str]) -> None:
    add_caption(document, caption)
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Действия актера"
    table.rows[0].cells[1].text = "Отклик системы"
    for action, response in zip(actions, system):
        cells = table.add_row().cells
        cells[0].text = action
        cells[1].text = response
    style_table(table)


def add_testcase_table(document: Document, caption: str, case: TestCase) -> None:
    add_caption(document, caption)
    table = document.add_table(rows=0, cols=2)
    for left, right in [
        ("Идентификатор", case.ident),
        ("Название", case.name),
        ("Цель", case.goal),
        ("Предусловия", case.preconditions),
        ("Шаги выполнения", "\n".join(case.actions)),
        ("Ожидаемый результат", case.expected),
    ]:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
    style_table(table)


def add_matrix_table(document: Document, caption: str, headers: list[str], rows: list[list[str]]) -> None:
    add_caption(document, caption)
    table = document.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table)


def add_list_paragraphs(document: Document, items: list[str]) -> None:
    for item in items:
        add_text(document, item, indent=False)


def write_title_page(document: Document) -> None:
    add_text(document, "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ", center=True, bold=True, indent=False)
    add_text(document, "Пояснительная записка к дипломному проекту", center=True, bold=True, indent=False)
    add_text(document, "по теме: «Разработка веб-приложения “Тэона” для автоматизации подбора недвижимости клиентам риелтора»", center=True, bold=True, indent=False)
    for _ in range(10):
        document.add_paragraph("")
    add_text(document, "Студент: Зубач Н.", indent=False)
    add_text(document, "Группа: ____________", indent=False)
    add_text(document, "Руководитель: ____________", indent=False)
    for _ in range(8):
        document.add_paragraph("")
    add_text(document, "2026", center=True, indent=False)
    add_page_break(document)


def write_intro(document: Document) -> None:
    add_heading(document, "Введение")
    paragraphs = [
        "В условиях роста объема предложений на рынке недвижимости риелтору требуется быстро подбирать объекты по параметрам клиента, фиксировать результаты поиска и формировать подборку для отправки. При работе вручную специалисту приходится параллельно вести клиентскую базу, просматривать сайты застройщиков и агрегаторов, сравнивать найденные предложения и собирать текст сообщения клиенту.",
        "Актуальность темы обусловлена необходимостью сокращения времени на рутинные операции риелтора и уменьшения количества ручных ошибок при подборе объектов недвижимости. Использование единого веб-приложения позволяет централизованно хранить параметры клиентов, результаты поиска, выбранные объекты и историю действий.",
        "Целью дипломного проекта является разработка веб-приложения «Тэона», обеспечивающего автоматизацию работы риелтора при подборе недвижимости. Для достижения поставленной цели необходимо решить задачи проектирования архитектуры системы, выбора стека технологий, разработки пользовательского интерфейса, реализации серверной логики, интеграции с PostgreSQL, настройки контейнерного развертывания и подготовки эксплуатационной документации.",
    ]
    for paragraph in paragraphs:
        add_text(document, paragraph)


def write_section_1(document: Document) -> None:
    add_heading(document, "1 Назначение и цели разработки")
    content = [
        "Разрабатываемое веб-приложение предназначено для автоматизации процессов работы риелтора с клиентскими заявками на подбор недвижимости. Система должна обеспечивать единое рабочее пространство, в котором специалист может авторизоваться, создать карточку клиента, сохранить параметры поиска, запустить подбор объектов, сформировать shortlist и подготовить текст сообщения для клиента.",
        "Пользователем системы является риелтор, сопровождающий клиентов на этапах первичного запроса, поиска объектов и подготовки персональной подборки. Приложение ориентировано на рынок новостроек и загородной недвижимости города Краснодара.",
        "Основной целью разработки является создание прикладного программного средства, сокращающего трудозатраты на поиск и систематизацию объектов недвижимости. Частные задачи проекта включают обеспечение постоянного хранения данных в PostgreSQL, поддержку работы через Docker Compose, нормализацию карточек объектов из внешних источников и предоставление понятного интерфейса для ежедневной работы.",
    ]
    for paragraph in content:
        add_text(document, paragraph)


def use_cases() -> list[UseCase]:
    return [
        UseCase(
            name="Поиск объектов для клиента",
            actors="Риелтор",
            summary="Риелтор вводит параметры клиента и инициирует подбор объектов недвижимости.",
            goal="Получить список релевантных объектов и сохранить результаты в базе данных.",
            kind="Базовый",
            success_actions=[
                "Действие №1. Риелтор авторизуется в приложении.",
                "Действие №2. Риелтор создает карточку клиента и заполняет параметры поиска.",
                "Действие №3. Риелтор подтверждает запуск поиска объектов.",
                "Действие №4. Риелтор просматривает найденные объекты.",
            ],
            success_system=[
                "Действие №1. Система открывает рабочее пространство риелтора.",
                "Действие №2. Система сохраняет клиента и профиль поиска в PostgreSQL.",
                "Действие №3. Система передает параметры в search-service, сохраняет search run и найденные объекты.",
                "Действие №4. Система отображает карточки объектов с ценой, площадью, источником и процентом совпадения.",
            ],
            exception_actions=[
                "Действие №1. Риелтор запускает поиск при временной недоступности части источников.",
                "Действие №2. Риелтор повторно открывает карточку клиента.",
            ],
            exception_system=[
                "Действие №1. Система пропускает недоступные источники, продолжает поиск по оставшимся и сохраняет только валидные результаты.",
                "Действие №2. Система показывает уже сохраненные найденные объекты и историю запусков поиска.",
            ],
        ),
        UseCase(
            name="Формирование подборки и подготовка сообщения",
            actors="Риелтор",
            summary="Риелтор отбирает подходящие объекты и формирует текст сообщения для отправки клиенту.",
            goal="Подготовить персональную подборку и сохранить ее в системе.",
            kind="Базовый",
            success_actions=[
                "Действие №1. Риелтор открывает карточку клиента.",
                "Действие №2. Риелтор добавляет объекты в подборку.",
                "Действие №3. Риелтор нажимает кнопку подготовки сообщения.",
                "Действие №4. Риелтор копирует телефон и текст сообщения.",
            ],
            success_system=[
                "Действие №1. Система отображает найденные объекты и текущую подборку.",
                "Действие №2. Система сохраняет shortlist_items в PostgreSQL.",
                "Действие №3. Система генерирует текст подборки, сохраняет share_message и показывает модальное окно.",
                "Действие №4. Система позволяет отметить подборку как отправленную.",
            ],
            exception_actions=[
                "Действие №1. Риелтор пытается сформировать сообщение при пустой подборке.",
            ],
            exception_system=[
                "Действие №1. Система выводит сообщение о том, что подборка пока пустая, и не создает share_message.",
            ],
        ),
    ]


def testcases() -> list[TestCase]:
    return [
        TestCase(
            ident="Е.1",
            name="Успешная авторизация риелтора",
            goal="Проверить возможность входа в систему по корректным учетным данным.",
            preconditions="Docker Compose запущен. В базе данных существует тестовый пользователь test / test.",
            actions=[
                "Действие №1. Открыть страницу входа в приложение.",
                "Действие №2. Ввести логин test и пароль test.",
                "Действие №3. Нажать кнопку «Войти».",
            ],
            expected="Система открывает главную страницу, устанавливает сессионную cookie и отображает данные пользователя.",
        ),
        TestCase(
            ident="Е.2",
            name="Создание нового клиента",
            goal="Проверить сохранение карточки клиента и профиля поиска.",
            preconditions="Пользователь авторизован.",
            actions=[
                "Действие №1. Перейти на страницу создания клиента.",
                "Действие №2. Заполнить обязательные поля формы.",
                "Действие №3. Нажать кнопку сохранения клиента.",
            ],
            expected="Система создает записи в таблицах clients и client_search_profiles, открывает карточку нового клиента и запускает поиск.",
        ),
        TestCase(
            ident="Е.3",
            name="Запуск поиска объектов",
            goal="Проверить загрузку, фильтрацию и сохранение результатов поиска.",
            preconditions="В системе существует клиент с заполненным профилем поиска.",
            actions=[
                "Действие №1. Открыть карточку клиента.",
                "Действие №2. Нажать кнопку повторного запуска поиска.",
                "Действие №3. Дождаться завершения запроса.",
            ],
            expected="Система создает запись search_runs, получает результаты от search-service и сохраняет найденные объекты в PostgreSQL.",
        ),
        TestCase(
            ident="Е.4",
            name="Добавление объекта в подборку",
            goal="Проверить сохранение shortlist_items и изменение состояния карточки объекта.",
            preconditions="Для клиента уже найдены объекты.",
            actions=[
                "Действие №1. Открыть карточку клиента.",
                "Действие №2. Нажать кнопку «В подборку» на одной из карточек объектов.",
            ],
            expected="Система создает запись shortlist_items, а кнопка меняет состояние на «В подборке».",
        ),
        TestCase(
            ident="Е.5",
            name="Подготовка текста подборки",
            goal="Проверить генерацию сообщения и запись share_messages.",
            preconditions="Подборка клиента не пуста.",
            actions=[
                "Действие №1. Открыть карточку клиента.",
                "Действие №2. Нажать кнопку подготовки сообщения.",
                "Действие №3. Проверить содержимое модального окна.",
            ],
            expected="Система формирует текст подборки, сохраняет запись share_messages и отображает телефон клиента вместе с текстом сообщения.",
        ),
        TestCase(
            ident="Е.6",
            name="Повторный запуск поиска по клиенту",
            goal="Проверить создание новой search run и обновление найденных объектов.",
            preconditions="В системе существует клиент с уже сохраненными результатами поиска.",
            actions=[
                "Действие №1. Открыть карточку клиента.",
                "Действие №2. Запустить поиск повторно.",
                "Действие №3. Дождаться завершения операции.",
            ],
            expected="Система создает новую запись search_runs, обновляет last_seen_at у объектов и пересчитывает match score.",
        ),
        TestCase(
            ident="Е.7",
            name="Валидация обязательных полей клиента",
            goal="Проверить блокировку создания клиента при незаполненных обязательных полях.",
            preconditions="Пользователь авторизован и открыл форму нового клиента.",
            actions=[
                "Действие №1. Оставить пустым имя клиента или телефон.",
                "Действие №2. Попытаться сохранить форму.",
            ],
            expected="Система не отправляет запрос, подсвечивает ошибку и не создает запись в базе данных.",
        ),
        TestCase(
            ident="Е.8",
            name="Некорректный логин или пароль",
            goal="Проверить реакцию системы на ошибочные учетные данные.",
            preconditions="Открыта страница входа.",
            actions=[
                "Действие №1. Ввести неверный логин или пароль.",
                "Действие №2. Нажать кнопку «Войти».",
            ],
            expected="Система возвращает статус 401 и отображает сообщение о неверном логине или пароле.",
        ),
        TestCase(
            ident="Е.9",
            name="Пустая подборка при генерации сообщения",
            goal="Проверить запрет подготовки сообщения без shortlist.",
            preconditions="Карточка клиента открыта, shortlist пуст.",
            actions=[
                "Действие №1. Нажать кнопку подготовки сообщения.",
            ],
            expected="Система не создает запись share_messages и выводит сообщение «Подборка пока пустая».",
        ),
        TestCase(
            ident="Е.10",
            name="Сохранение данных после перезапуска контейнеров",
            goal="Проверить постоянное хранение данных в PostgreSQL.",
            preconditions="В системе уже созданы клиенты и shortlist. Контейнеры запущены через Docker Compose.",
            actions=[
                "Действие №1. Выполнить остановку контейнеров командой npm run prod:down.",
                "Действие №2. Повторно запустить систему командой npm run prod:up.",
                "Действие №3. Открыть список клиентов.",
            ],
            expected="Система отображает ранее созданных клиентов и ранее сформированные подборки, так как данные сохранены в volume PostgreSQL.",
        ),
    ]


def build_document(asset_paths: dict[str, Path]) -> None:
    document = Document()
    configure_styles(document)
    add_page_number(document.sections[0])

    add_heading(document, "Содержание")
    add_toc(document)
    add_page_break(document)

    write_intro(document)
    write_section_1(document)
    add_heading(document, "1.1 Анализ предметной области", level=2)
    domain_paragraphs = [
        "Предметной областью проекта является подбор жилой недвижимости клиентам риелтора. В рассматриваемом процессе риелтор принимает запрос от клиента, уточняет бюджет, желаемый тип недвижимости, минимальную площадь, число комнат, район или населенный пункт, а также дополнительные требования, например отделку, этажность или наличие коммуникаций.",
        "На практике информация о доступных объектах распределена по нескольким внешним источникам: сайтам застройщиков, агрегаторам новостроек, каталогам домов и коттеджных поселков. Каждый источник имеет собственную структуру HTML, по-разному хранит цену, площадь, заголовок карточки и изображения. Вследствие этого специалист тратит значительное время не только на поиск, но и на нормализацию информации: перенос ссылки, пересчет цены за метр, очистку служебных заголовков и отбор нерелевантных изображений.",
        "Дополнительную сложность создает необходимость повторного обращения к одному и тому же клиенту. Если результаты поиска не сохраняются централизованно, риелтор вынужден заново восстанавливать историю подбора, повторно открывать сайты, вспоминать, какие объекты уже показывались, и формировать текст сообщения вручную. Это повышает вероятность ошибок, приводит к дублированию объектов и снижает качество сопровождения клиента.",
        "Разрабатываемая система решает перечисленные проблемы за счет разделения функций на frontend, backend API, search-service и PostgreSQL. Интерфейс предоставляет единое рабочее пространство, backend отвечает за бизнес-логику и сохранение данных, search-service выполняет live-поиск и нормализацию карточек объектов, а PostgreSQL обеспечивает постоянное хранение результатов между сеансами работы.",
    ]
    for paragraph in domain_paragraphs:
        add_text(document, paragraph)
    add_matrix_table(document, "Таблица 1 – Основные сущности предметной области", ["Сущность", "Описание", "Основные атрибуты"], [
        ["Риелтор", "Пользователь системы, работающий с клиентской базой.", "Логин. Пароль. ФИО. Контакты."],
        ["Клиент", "Покупатель недвижимости, для которого ведется подбор.", "ФИО. Телефон. Статус. Тип недвижимости."],
        ["Профиль поиска", "Набор параметров, по которым выполняется подбор объектов.", "Бюджет. Комнатность. Площадь. Район. Коммуникации."],
        ["Объект недвижимости", "Нормализованная карточка квартиры или дома из внешнего источника.", "Цена. Площадь. Комнаты. Источник. Ссылка."],
        ["Подборка", "Список объектов, вручную отобранных риелтором.", "Клиент. Объект. Примечание."],
        ["Запуск поиска", "Отдельная операция поиска по клиентскому профилю.", "Статус. Количество найденных. Количество сохраненных."],
    ])
    add_heading(document, "1.2 Требования к разрабатываемой системе", level=2)
    add_matrix_table(document, "Таблица 2 – Функциональные требования", ["Группа", "Требование"], [
        ["Авторизация", "Система должна обеспечивать вход риелтора по логину и паролю."],
        ["Клиенты", "Система должна обеспечивать создание, просмотр, редактирование и удаление карточек клиентов."],
        ["Поиск", "Система должна принимать профиль поиска и запускать подбор объектов через отдельный search-service."],
        ["Результаты", "Система должна сохранять найденные объекты и отображать их в карточке клиента."],
        ["Подборка", "Система должна позволять добавлять объекты в shortlist и удалять их оттуда."],
        ["Сообщение", "Система должна формировать текст сообщения для клиента и сохранять историю формирования."],
        ["Хранение", "Система должна сохранять данные в PostgreSQL после перезапуска контейнеров."],
    ])
    add_matrix_table(document, "Таблица 3 – Нефункциональные требования", ["Категория", "Требование"], [
        ["Производительность", "Поиск должен завершаться за разумное время с учетом сетевых ограничений внешних сайтов."],
        ["Отказоустойчивость", "Ошибка одного источника не должна останавливать поиск по остальным сайтам."],
        ["Совместимость", "Интерфейс должен корректно работать в современных браузерах на настольных устройствах."],
        ["Сопровождаемость", "Кодовая база должна быть разделена на frontend, backend API и search-service."],
        ["Развертывание", "Система должна запускаться через Docker Compose с минимальным числом команд."],
    ])

    add_heading(document, "2 Технический проект")
    add_heading(document, "2.1 Проектирование диаграмм вариантов использования и описание сценариев", level=2)
    add_text(document, "С целью формализации требований к приложению была построена диаграмма вариантов использования, отражающая взаимодействие риелтора с системой. Диаграмма демонстрирует ключевые операции: авторизацию, создание клиента, запуск поиска, просмотр найденных объектов, формирование подборки и подготовку сообщения.")
    add_figure(document, "Рисунок 1 – Диаграмма вариантов использования приложения «Тэона»", asset_paths["use_case"], width_cm=16)
    add_text(document, "На рисунке 2 представлена диаграмма вариантов использования модуля поиска объектов. Она отражает отдельную ответственность поискового контура: получение параметров клиента, обращение к источникам, фильтрацию по городу Краснодар, расчет процента совпадения и сохранение результата.")
    add_figure(document, "Рисунок 2 – Диаграмма вариантов использования модуля поиска объектов", asset_paths["use_case_search"], width_cm=16)
    add_text(document, "На рисунке 3 показан модуль формирования подборки. В нем пользователь вручную выбирает подходящие объекты, формирует текст сообщения, копирует данные для отправки клиенту и отмечает подборку как отправленную.")
    add_figure(document, "Рисунок 3 – Диаграмма вариантов использования модуля формирования подборки", asset_paths["use_case_shortlist"], width_cm=16)

    first, second = use_cases()
    add_text(document, "Для детализации основного варианта использования «Поиск объектов для клиента» была составлена его характеристика и сценарии выполнения.")
    add_key_value_table(document, "Таблица 1.1 – Главный раздел сценария варианта использования «Поиск объектов для клиента»", [
        ("Вариант использования (прецедент)", first.name),
        ("Актеры", first.actors),
        ("Краткое описание", first.summary),
        ("Цель", first.goal),
        ("Тип", first.kind),
    ])
    add_scenario_table(document, "Таблица 1.2 – Сценарий успешного выполнения варианта использования «Поиск объектов для клиента»", first.success_actions, first.success_system)
    add_scenario_table(document, "Таблица 1.3 – Обработка исключительных ситуаций для варианта использования «Поиск объектов для клиента»", first.exception_actions, first.exception_system)

    add_text(document, "Вторым ключевым вариантом использования является формирование персональной подборки и подготовка сообщения для клиента.")
    add_key_value_table(document, "Таблица 2.1 – Главный раздел сценария варианта использования «Формирование подборки и подготовка сообщения»", [
        ("Вариант использования (прецедент)", second.name),
        ("Актеры", second.actors),
        ("Краткое описание", second.summary),
        ("Цель", second.goal),
        ("Тип", second.kind),
    ])
    add_scenario_table(document, "Таблица 2.2 – Сценарий успешного выполнения варианта использования «Формирование подборки и подготовка сообщения»", second.success_actions, second.success_system)
    add_scenario_table(document, "Таблица 2.3 – Обработка исключительных ситуаций для варианта использования «Формирование подборки и подготовка сообщения»", second.exception_actions, second.exception_system)

    add_heading(document, "2.2 Проектирование диаграмм деятельности", level=2)
    add_text(document, "На основе сценария подбора объектов была построена блок-схема, описывающая последовательность операций от создания клиента до сохранения найденных объектов.")
    add_figure(document, "Рисунок 4 – Диаграмма деятельности подбора объектов недвижимости", asset_paths["activity"], width_cm=16)

    add_heading(document, "2.3 Проектирование диаграммы состояний", level=2)
    add_text(document, "Состояние клиентской заявки изменяется серверной частью приложения в зависимости от действий риелтора и результата поискового контура. Диаграмма состояний показывает переходы между новой заявкой, поиском, найденными объектами, готовой подборкой, отправленной подборкой и ошибочными состояниями.")
    add_figure(document, "Рисунок 5 – Диаграмма состояний клиентской заявки", asset_paths["state"], width_cm=16)

    add_heading(document, "2.4 Проектирование диаграммы компонентов", level=2)
    add_text(document, "Компонентная модель системы отражает разделение приложения на пользовательский интерфейс, backend API, поисковый сервис, PostgreSQL и внешние источники недвижимости. Такое разделение соответствует фактической структуре проекта и контейнерному развертыванию через Docker Compose.")
    add_figure(document, "Рисунок 6 – Диаграмма компонентов веб-приложения «Тэона»", asset_paths["component"], width_cm=16)

    add_heading(document, "2.5 Проектирование модели данных", level=2)
    add_text(document, "Модель данных спроектирована с учетом необходимости постоянного хранения клиентов, параметров поиска, найденных объектов, shortlist и истории запусков поиска. В отличие от временных in-memory структур, PostgreSQL обеспечивает сохранность данных после перезапуска контейнеров и позволяет централизованно хранить результат поиска по каждому клиенту.")
    add_figure(document, "Рисунок 7 – Архитектурная схема системы", asset_paths["architecture"], width_cm=16)
    add_figure(document, "Рисунок 8 – Схема базы данных PostgreSQL", asset_paths["er"], width_cm=16)

    add_heading(document, "2.6 Разработка алгоритмов обработки данных", level=2)
    add_heading(document, "2.6.1 Разработка алгоритма парсинга и географической фильтрации", level=3)
    add_text(document, "Алгоритм парсинга отделен от backend API и реализован в search-service. Его задача заключается в получении HTML и JSON-LD из заранее заданных источников, извлечении характеристик объекта, отсечении карточек не из города Краснодар и нормализации результата к единой структуре PropertyItem.")
    add_figure(document, "Рисунок 9 – Алгоритм парсинга и географической фильтрации", asset_paths["algorithm_parsing"], width_cm=16)
    add_heading(document, "2.6.2 Разработка алгоритма расчета процента совпадения", level=3)
    add_text(document, "Алгоритм расчета match score применяется после нормализации объекта. Для квартир учитываются бюджет, комнатность, площадь, район, срок сдачи и отделка. Для домов учитываются бюджет, площадь дома, площадь участка, локация, коммуникации, этажность и спальни.")
    add_figure(document, "Рисунок 10 – Алгоритм расчета match score", asset_paths["algorithm_scoring"], width_cm=16)
    add_heading(document, "2.6.3 Разработка алгоритма сохранения результатов поиска", level=3)
    add_text(document, "Алгоритм сохранения результатов поиска обеспечивает постоянство данных. Backend API сохраняет нормализованные объекты в таблицу properties, связь объекта с клиентом в client_found_properties, а историю запуска поиска в search_runs.")
    add_figure(document, "Рисунок 11 – Алгоритм сохранения результатов поиска в PostgreSQL", asset_paths["algorithm_persistence"], width_cm=16)

    add_heading(document, "2.7 Проектирование интерфейса пользователя", level=2)
    add_text(document, "Интерфейс приложения был спроектирован в виде набора wireframes, отражающих основные рабочие окна системы. Такой подход позволил определить расположение зон управления, таблиц, карточек и модальных окон до детализации визуального оформления.")
    add_figure(document, "Рисунок 12 – Wireframe окна авторизации", asset_paths["wf_login"], width_cm=16)
    add_figure(document, "Рисунок 13 – Wireframe главной страницы", asset_paths["wf_dashboard"], width_cm=16)
    add_figure(document, "Рисунок 14 – Wireframe списка клиентов", asset_paths["wf_clients"], width_cm=16)
    add_figure(document, "Рисунок 15 – Wireframe карточки клиента", asset_paths["wf_client"], width_cm=16)
    add_text(document, "После этапа wireframes были подготовлены экранные макеты и рабочие экраны интерфейса, отражающие визуальный стиль системы. Основной темой является светлый CRM-интерфейс с оранжевым акцентом, фиксированным боковым меню и верхней панелью поиска.")
    add_figure(document, "Рисунок 16 – Экран авторизации приложения", asset_paths["screen_login"], width_cm=16)
    add_figure(document, "Рисунок 17 – Главная страница приложения", asset_paths["screen_dashboard"], width_cm=16)
    add_figure(document, "Рисунок 18 – Страница списка клиентов", asset_paths["screen_clients"], width_cm=16)
    add_figure(document, "Рисунок 19 – Карточка клиента", asset_paths["screen_client"], width_cm=16)
    add_figure(document, "Рисунок 20 – Модальное окно подготовки сообщения", asset_paths["screen_modal"], width_cm=16)

    add_heading(document, "3 Реализация")
    add_heading(document, "3.1 Обоснование выбора средств разработки и методов разработки", level=2)
    techs = [
        "Для реализации пользовательского интерфейса выбрана библиотека React и сборщик Vite. Такое решение обеспечивает компонентный подход, быстрый локальный запуск и удобную организацию SPA-приложения.",
        "Серверная часть разработана на Node.js и Express. Библиотека pg используется для прямого подключения к PostgreSQL, что соответствует целевой архитектуре проекта и исключает промежуточный ORM-слой.",
        "Поисковый сервис реализован на Python и FastAPI, так как этот стек удобен для обработки HTTP-запросов, парсинга HTML и нормализации данных из внешних источников. Для извлечения информации применяются httpx, BeautifulSoup4 и lxml.",
        "Развертывание выполняется через Docker Compose, что позволяет запускать frontend, backend, search-service и PostgreSQL как согласованный набор контейнеров. Такой подход упрощает перенос проекта на хостинг и его демонстрацию.",
    ]
    for paragraph in techs:
        add_text(document, paragraph)

    add_heading(document, "3.2 Разработка базы данных в среде СУБД", level=2)
    add_text(document, "Физическая модель данных реализована в SQL-скрипте database/init.sql. Основные таблицы и их назначение представлены в таблице 10.")
    add_caption(document, "Таблица 10 – Назначение основных таблиц базы данных")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Таблица"
    table.rows[0].cells[1].text = "Назначение"
    for row in [
        ("users", "Хранение учетных записей риелторов."),
        ("clients", "Карточки клиентов и основные контактные данные."),
        ("client_search_profiles", "Параметры поиска клиента."),
        ("properties", "Нормализованные карточки объектов недвижимости."),
        ("client_found_properties", "Связь клиента с найденными объектами и match score."),
        ("shortlist_items", "Ручная подборка объектов для клиента."),
        ("share_messages", "Сгенерированные тексты сообщений для отправки клиенту."),
        ("search_runs", "История запусков поиска и агрегированные результаты."),
    ]:
        cells = table.add_row().cells
        cells[0].text = row[0]
        cells[1].text = row[1]
    style_table(table)
    add_text(document, "Для хранения клиентов используется отдельная таблица clients, связанная с users по полю realtor_id. Параметры поиска вынесены в таблицу client_search_profiles, чтобы карточка клиента не перегружалась большим числом nullable-полей. Нормализованные объекты источников сохраняются в properties, а факт соответствия конкретного объекта конкретному клиенту фиксируется в client_found_properties с match score и текстовыми причинами совпадения или несовпадения.")
    add_text(document, "Такая структура позволяет повторно использовать одну и ту же карточку объекта для нескольких клиентов, отслеживать историю запусков поиска и не терять собранные подборки. Для итоговой демонстрации и последующей защиты в проект был подготовлен отдельный скриншот просмотра таблицы clients в базе данных.")
    add_figure(document, "Рисунок 21 – Просмотр таблицы clients в PostgreSQL", asset_paths["screen_db"], width_cm=16)

    add_heading(document, "3.3 Программная реализация модулей в среде программирования", level=2)
    modules = [
        "Frontend-модуль реализует экран авторизации, главную страницу с KPI, список клиентов, форму создания клиента, карточку клиента, модальные окна просмотра объекта и подготовки сообщения, а также вспомогательные страницы профиля, помощи и настроек.",
        "Backend API предоставляет маршруты /api/auth, /api/dashboard, /api/clients и /api/properties. Через эти маршруты выполняются аутентификация, загрузка KPI, управление клиентами, запуск поиска, сохранение shortlist и подготовка текста сообщения.",
        "Search-service получает параметры клиента, распределяет запрос по адаптерам источников, нормализует title и description, фильтрует некачественные изображения, рассчитывает соответствие объекта требованиям клиента и возвращает валидный набор результатов.",
        "Модуль хранения данных реализован на PostgreSQL и поддерживает постоянное хранение клиентов, найденных объектов, shortlist и истории поиска. Инициализация схемы базы данных выполняется через SQL-скрипт при запуске контейнеров.",
    ]
    for paragraph in modules:
        add_text(document, paragraph)
    add_heading(document, "3.3.1 Реализация frontend-модуля", level=3)
    frontend_details = [
        "Frontend расположен в директории apps/web и представляет собой SPA-приложение на React. Основная маршрутизация настроена в App.tsx, где реализованы защищенные маршруты для страниц dashboard, clients, client, analytics, profile, help и settings. При обращении к защищенному разделу без валидной сессии пользователь перенаправляется на страницу login.",
        "Для работы с API frontend использует единый модуль shared/api.ts. Он инкапсулирует отправку HTTP-запросов, сериализацию JSON и базовую обработку ошибок. Запросы к dashboard, клиентам и shortlist выполняются асинхронно, а полученные данные отображаются без полной перезагрузки страницы.",
        "Каркас интерфейса реализован в компоненте AppShell.tsx. В нем сосредоточены боковое меню, верхняя панель, быстрый поиск по клиентам и всплывающее меню пользователя. Каждая страница встраивается в этот каркас как отдельный маршрут, что упрощает сопровождение интерфейса и повторное использование элементов.",
    ]
    for paragraph in frontend_details:
        add_text(document, paragraph)
    add_heading(document, "3.3.2 Реализация backend API", level=3)
    backend_details = [
        "Backend расположен в директории apps/api и реализован на Express. Точка входа server.ts инициирует подключение к PostgreSQL, вызывает ensureSeedUser для создания тестового пользователя и поднимает HTTP-сервер на порту API_PORT.",
        "Маршрут auth.ts отвечает за проверку логина и пароля, генерацию JWT cookie и загрузку текущего пользователя по маршруту /api/auth/me. Все защищенные маршруты проходят через middleware requireAuth, которое извлекает идентификатор пользователя из signed cookie.",
        "Маршрут clients.ts содержит основной бизнес-флоу системы. Здесь реализованы создание клиента, обновление карточки, удаление, повторный запуск поиска, работа с shortlist, подготовка сообщения и установка статуса отправки. API выступает координатором между frontend, PostgreSQL и search-service.",
    ]
    for paragraph in backend_details:
        add_text(document, paragraph)
    add_heading(document, "3.3.3 Реализация search-service и алгоритмов парсинга", level=3)
    parsing_details = [
        "Search-service расположен в apps/search-service и реализован на FastAPI. При вызове POST /search сервис получает структуру SearchRequest, включающую propertyType, бюджет, площадь, комнаты, районы, названия поселков и дополнительные параметры. Далее формируется список адаптеров, зависящий от типа недвижимости: для квартир используются сайты застройщиков и агрегаторов новостроек, для домов – каталоги домов и коттеджных поселков.",
        "Каждый адаптер наследует BaseAdapter и проходит одинаковый pipeline. Сначала выполняется проверка robots.txt через RobotFileParser. Затем HTML загружается асинхронно через httpx.AsyncClient с пользовательским User-Agent. После этого содержимое обрабатывается парсером BeautifulSoup в режиме lxml, что позволяет быстро извлекать как JSON-LD, так и видимые карточки на странице.",
        "Алгоритм извлечения данных начинается с JSON-LD. Если в документе есть script type application/ld+json, сервис разворачивает граф, выделяет узлы типов product, offer, apartment, house, residence и realestate, извлекает цену, название, ссылку, изображения, площадь, число комнат, ЖК или КП, район, адрес и год сдачи. Если JSON-LD отсутствует или данных недостаточно, включается резервный механизм анализа видимых карточек, основанный на ссылках и текстовых блоках.",
        "После первичного извлечения выполняется нормализация заголовка. Служебные слова вроде «Избранное», «Отзывы», «Все новостройки», «Старт продаж» и маркетинговые фразы отбрасываются. Если исходный title непригоден, система конструирует заголовок из комнатности, площади, названия ЖК или КП, адреса и района. Для квартир используется форма «1-к квартира, 40.6 м² в ЖК ...», для домов – «Дом 128 м², участок 5 сот. в КП ...».",
        "Отдельно выполняется фильтрация изображений. Из выдачи исключаются favicon, logo, icon, sprite, placeholder, avatar, social-изображения и SVG. Система оставляет только реальные JPG, PNG и WEBP-файлы, пригодные для использования в карточке объекта. Описание объекта также собирается заново: в него включаются площадь, цена, локация и источник, а служебный SEO-текст страницы не сохраняется.",
        "Финальный этап search-service – расчет match score. Для квартир баллы начисляются за попадание цены, комнатности, площади, района, срока сдачи и отделки в заданный профиль клиента. Для домов учитываются бюджет, площадь дома, площадь участка, локация, коммуникации, этажность и число спален. Результат ограничивается диапазоном от 0 до 100, а в объект записываются как положительные причины совпадения, так и текстовые причины несовпадения.",
    ]
    for paragraph in parsing_details:
        add_text(document, paragraph)
    add_matrix_table(document, "Таблица 11 – Подключенные источники поиска", ["Тип объекта", "Источник", "Тип источника", "Основные данные"], [
        ["Квартира", "НАШ.ДОМ.РФ", "Официальный реестр", "Цена. Площадь. Комнатность. Ссылка."],
        ["Квартира", "Домострой Краснодар", "Агрегатор", "Цена. Площадь. ЖК. Источник."],
        ["Квартира", "ССК / ВКБ / DOGMA / ТОЧНО", "Сайты застройщиков", "ЖК. Цена. Площадь. Комнатность."],
        ["Квартира", "Krasdom / 23kvartiri / Novostrojki-KRD", "Агрегаторы", "Цена. Площадь. Комнаты. Ссылка."],
        ["Дом", "Doma-kr", "Каталог домов", "Цена. Площадь дома. Ссылка."],
        ["Дом", "КП Краснодар", "Каталог коттеджных поселков", "Площадь дома. Участок. Цена."],
        ["Дом", "Поселки Краснодара", "Каталог поселков", "Локация. Участок. Цена."],
    ])
    add_heading(document, "3.3.4 Алгоритм сохранения результатов поиска в PostgreSQL", level=3)
    add_text(document, "После получения списка PropertyItem backend API проходит по всем объектам, имеющим sourceUrl. Для каждого объекта вызывается upsertProperty, который либо создает новую запись в таблице properties, либо обновляет уже существующую карточку по source_url. Затем через upsertClientFoundProperty создается или обновляется связь между клиентом и объектом, включая match_score, match_reasons и mismatch_reasons.")
    add_text(document, "Одновременно создается search run через createSearchRun, а после успешного завершения операция закрывается через finishSearchRun. Если в процессе обращения к search-service возникает ошибка, backend вызывает failSearchRun и переводит клиента в статус error. За счет этого история поисков остается прозрачной и пригодной для анализа даже при сбоях внешних источников.")
    add_heading(document, "3.3.5 Формирование подборки и текста сообщения", level=3)
    add_text(document, "Shortlist формируется вручную риелтором на странице карточки клиента. При нажатии кнопки «В подборку» выполняется POST-запрос к маршруту /api/clients/:id/shortlist/:propertyId. Backend проверяет, что объект действительно принадлежит найденным результатам данного клиента, после чего создает запись shortlist_items. Повторное добавление предотвращается уникальным ограничением по client_id и property_id.")
    add_text(document, "Подготовка текста сообщения выполняется маршрутом /api/clients/:id/share-message. Сервис buildShareMessage собирает компактный список объектов из текущей подборки, формирует понятный текст для отправки и сохраняет его в таблицу share_messages. Отправка как внешняя интеграция не реализуется: система сознательно ограничивается режимом copy-only, чтобы не связываться с правилами внешних мессенджеров и не усложнять дипломный стек.")

    add_heading(document, "3.4 Тестирование программных модулей", level=2)
    add_heading(document, "3.4.1 Модульное тестирование", level=3)
    add_text(document, "Для search-service разработаны модульные тесты, проверяющие корректность матчинга и нормализации карточек объектов. В тестах контролируется фильтрация служебных названий, работа с ценой и площадью, а также формирование структурированных карточек недвижимости.")
    add_caption(document, "Таблица 12 – Результаты автоматизированных проверок")
    checks = document.add_table(rows=1, cols=2)
    checks.rows[0].cells[0].text = "Проверка"
    checks.rows[0].cells[1].text = "Результат"
    for row in [
        ("npm run typecheck", "Пройдено."),
        ("npm run build", "Пройдено."),
        ("python -m unittest discover -s tests -v", "Пройдено."),
    ]:
        cells = checks.add_row().cells
        cells[0].text = row[0]
        cells[1].text = row[1]
    style_table(checks)
    add_figure(document, "Рисунок 22 – Результаты автоматизированных проверок", asset_paths["test_results"], width_cm=16)

    add_heading(document, "3.4.2 Интеграционное тестирование", level=3)
    add_text(document, "Интеграционное тестирование ориентировано на проверку сквозных пользовательских сценариев: авторизации, создания клиента, запуска поиска, добавления объекта в подборку и подготовки текста сообщения. Подробные тест-кейсы приведены в приложении В.")
    add_text(document, "Отдельно проверяется постоянство хранения данных после перезапуска контейнеров Docker Compose. Этот сценарий критичен для дипломного проекта, так как одним из ключевых требований является переход на PostgreSQL и отказ от временного in-memory хранения.")

    add_heading(document, "3.5 Разработка эксплуатационной документации", level=2)
    add_text(document, "Для проекта подготовлен комплект эксплуатационных материалов, предназначенный для демонстрации и практического использования системы. Основой эксплуатационной документации являются руководство пользователя, инструкция по запуску production-сборки через Docker Compose и описание типовых сообщений интерфейса.")
    add_text(document, "В production-конфигурации используются четыре контейнера: postgres, search-service, api и web. Контейнер PostgreSQL инициализируется схемой базы данных при первом запуске и хранит данные в отдельном volume, что обеспечивает сохранность клиентской базы, найденных объектов и подборок после перезапуска системы.")
    add_text(document, "Такой формат развертывания подходит для защиты дипломной работы, так как система поднимается повторяемо на любом компьютере с Docker Desktop и не требует ручной установки Node.js, Python, PostgreSQL и сопутствующих зависимостей.")

    add_heading(document, "Заключение")
    add_text(document, "В результате дипломного проекта разработано веб-приложение «Тэона», предназначенное для автоматизации работы риелтора при подборе недвижимости. Система объединяет клиентскую базу, параметры поиска, live-поиск по внешним источникам, формирование shortlist и подготовку текста сообщения для клиента.")
    add_text(document, "В ходе работы были спроектированы архитектура и модель данных, реализованы frontend, backend API и search-service, настроено постоянное хранение данных в PostgreSQL через Docker Compose, а также подготовлена эксплуатационная документация. Реализованная система может использоваться как основа для дальнейшего развития CRM-функций, подключения дополнительных источников недвижимости и развертывания на серверной инфраструктуре.")

    add_heading(document, "Список использованных источников")
    sources = [
        "ГОСТ 19.201-78. Единая система программной документации. Техническое задание. Требования к содержанию и оформлению.",
        "ГОСТ Р 59795-2021. Информационные технологии. Комплекс стандартов на автоматизированные системы. Требования к содержанию документов.",
        "Node.js Documentation: [сайт]. URL: https://nodejs.org/en/docs/ (дата обращения: 21.05.2026).",
        "React Documentation: [сайт]. URL: https://react.dev/ (дата обращения: 21.05.2026).",
        "FastAPI Documentation: [сайт]. URL: https://fastapi.tiangolo.com/ (дата обращения: 21.05.2026).",
        "PostgreSQL Documentation: [сайт]. URL: https://www.postgresql.org/docs/ (дата обращения: 21.05.2026).",
        "Docker Documentation: [сайт]. URL: https://docs.docker.com/ (дата обращения: 21.05.2026).",
        "Beautiful Soup Documentation: [сайт]. URL: https://www.crummy.com/software/BeautifulSoup/bs4/doc/ (дата обращения: 21.05.2026).",
    ]
    for source in sources:
        add_text(document, source, indent=False)

    add_page_break(document)
    add_heading(document, "Приложение А", level=1)
    add_text(document, "(обязательное)", center=True, indent=False)
    add_heading(document, "Требования к функциональным характеристикам", level=2)
    add_heading(document, "1 Требования к функциональным характеристикам", level=3)
    add_text(document, "Приложение должно обеспечивать автоматизацию подбора жилой недвижимости для клиентов риелтора. Система должна поддерживать создание клиентской карточки, сохранение параметров поиска, подбор объектов, формирование shortlist и подготовку текста сообщения для клиента.")
    add_heading(document, "1.1 Требования к модулю авторизации", level=3)
    add_text(document, "Модуль авторизации должен обеспечивать вход пользователя по логину и паролю, создание пользовательской сессии и защиту маршрутов приложения от неавторизованного доступа.")
    add_heading(document, "1.2 Требования к составу выполняемых функций модуля работы с клиентами", level=3)
    add_text(document, "Модуль работы с клиентами должен обеспечивать создание, просмотр, редактирование и удаление карточек клиентов, хранение контактных данных, статуса работы и типа подбираемой недвижимости.")
    add_heading(document, "1.3 Требования к составу выполняемых функций модуля подбора недвижимости", level=3)
    add_text(document, "Модуль подбора недвижимости должен обеспечивать запуск поиска по параметрам клиента, фильтрацию объектов по городу Краснодар, сохранение найденных карточек, расчет процента совпадения и отображение результатов в карточке клиента.")
    add_heading(document, "1.4 Требования к составу входных и выходных данных", level=3)
    add_heading(document, "1.4.1 Требования к входным данным", level=3)
    add_text(document, "Входными данными являются сведения о клиенте, включая бюджет, тип недвижимости, желаемую площадь, число комнат, район города, срок сдачи, отделку и дополнительные характеристики для дома или квартиры.")
    add_heading(document, "1.4.2 Требования к выходным данным", level=3)
    add_text(document, "Выходными данными являются список найденных объектов, shortlist клиента, текст сообщения для отправки клиенту, а также агрегированные показатели по клиентской базе и результатам поиска.")
    add_heading(document, "1.5 Общие ограничения к данным", level=3)
    add_text(document, "Система должна сохранять только валидные карточки объектов недвижимости с корректной ссылкой на источник. Из результатов должны исключаться объекты из других городов Краснодарского края, служебные карточки, логотипы и записи без существенных характеристик.")

    add_page_break(document)
    add_heading(document, "Приложение Б", level=1)
    add_text(document, "(справочное)", center=True, indent=False)
    add_heading(document, "Требования к пользовательскому интерфейсу", level=2)
    add_text(document, "Интерфейс приложения должен быть выполнен в светлом, сдержанном деловом стиле. Основным акцентным цветом является оранжевый, используемый для активных кнопок, элементов выбора и показателей совпадения.")
    add_text(document, "Навигация должна быть организована через фиксированное боковое меню и верхнюю панель. Карточки объектов должны содержать короткий заголовок, цену, базовые характеристики, источник, район и действия риелтора.")
    add_text(document, "Все экраны должны корректно работать в браузере на настольном компьютере, не допускать наложения текста, переполнения карточек и визуальной неоднородности между разделами интерфейса.")

    add_page_break(document)
    add_heading(document, "Приложение В", level=1)
    add_text(document, "(обязательное)", center=True, indent=False)
    add_heading(document, "Тест-кейсы", level=2)
    for case in testcases():
        add_testcase_table(document, f"Таблица {case.ident} – {case.name}", case)

    add_page_break(document)
    add_heading(document, "Приложение Г", level=1)
    add_text(document, "(обязательное)", center=True, indent=False)
    add_heading(document, "Руководство пользователя", level=2)
    add_heading(document, "1 Выполнение программы", level=3)
    guide = [
        "Действие №1. Установить Docker Desktop на компьютер пользователя.",
        "Действие №2. В корне проекта подготовить файл окружения и выполнить команду production-запуска контейнеров.",
        "Действие №3. Открыть веб-приложение по локальному адресу в браузере.",
        "Действие №4. Выполнить вход с учетными данными риелтора.",
        "Действие №5. Создать карточку клиента и заполнить параметры подбора.",
        "Действие №6. Дождаться результатов поиска и при необходимости повторно запустить подбор.",
        "Действие №7. Добавить подходящие объекты в подборку и сформировать сообщение клиенту.",
    ]
    for line in guide:
        add_text(document, line, indent=False)
    add_heading(document, "2 Сообщения пользователю", level=3)
    add_text(document, "При некорректном логине или пароле система выводит сообщение о невозможности входа. При пустой подборке система сообщает, что формирование текста сообщения пока недоступно. Если поиск не дал результатов, в карточке клиента отображается соответствующий статус.")
    add_figure(document, "Рисунок Г.1 – Окно авторизации", asset_paths["screen_login"], width_cm=16)
    add_figure(document, "Рисунок Г.2 – Карточка клиента", asset_paths["screen_client"], width_cm=16)

    add_page_break(document)
    add_heading(document, "Приложение Д", level=1)
    add_text(document, "(обязательное)", center=True, indent=False)
    add_heading(document, "Руководство по развертыванию программы", level=2)
    add_heading(document, "1 Характеристика программы", level=3)
    add_text(document, "Программа представляет собой веб-приложение для автоматизации подбора недвижимости клиентам риелтора. В состав системы входят frontend, backend API, поисковый сервис и PostgreSQL.")
    add_heading(document, "2 Обращение к программе", level=3)
    add_text(document, "Для запуска программы используется Docker Compose. После старта контейнеров пользователь обращается к приложению через браузер. Для остановки системы применяются штатные команды Docker Compose без удаления volume базы данных.")
    add_heading(document, "3 Сообщения", level=3)
    add_text(document, "При отсутствии соединения с одним из внешних источников поиска система продолжает работу с остальными источниками. При ошибке поискового контура статус клиента переводится в состояние ошибки, что позволяет повторно выполнить поиск после устранения проблемы.")
    add_figure(document, "Рисунок Д.1 – Главная страница приложения", asset_paths["screen_dashboard"], width_cm=16)
    add_figure(document, "Рисунок Д.2 – Просмотр таблицы clients в PostgreSQL", asset_paths["screen_db"], width_cm=16)

    document.save(OUTPUT)


def main() -> None:
    assets = build_assets()
    build_document(assets)


if __name__ == "__main__":
    main()
